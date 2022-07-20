import os
import re
from typing import List

import numpy as np
import cv2 as cv
import pytesseract
from docx import Document
from pdf2image import convert_from_bytes


def read_from_pdf(pdf_file):
    """
    :param pdf_file: pdf file as binary file
    :return: returns a list of numpy array images
    """
    return list(map(np.array, convert_from_bytes(pdf_file, dpi=150)))


def get_cropped_parts(image_norm: np.array) -> List[np.array]:
    """
    :param image_norm: source image
    :return: list of crops
    """
    gray = cv.cvtColor(image_norm, cv.COLOR_BGR2GRAY)
    ret, thresh1 = cv.threshold(gray, 0, 255, cv.THRESH_OTSU | cv.THRESH_BINARY_INV)
    rect_kernel = cv.getStructuringElement(cv.MORPH_RECT, (18, 18))
    dilation = cv.dilate(thresh1, rect_kernel, iterations=1)
    contours, hierarchy = cv.findContours(dilation, cv.RETR_EXTERNAL, cv.CHAIN_APPROX_NONE)
    cropped_parts = []
    for cnt in contours:
        x, y, w, h = cv.boundingRect(cnt)
        cropped = image_norm[y:y + h, x:x + w]
        cropped_parts.append(cropped)
    return cropped_parts


def find_lines(img: np.array) -> np.array:
    """
    :param img: source image
    :return: handled image with lines contours
    """
    height, width = img.shape
    eroded = cv.erode(img, np.ones((3, 3)))
    kernel_length = 3 / 100
    horizontal_kernel = np.ones((1, int(width * kernel_length)))
    horizontal = cv.morphologyEx(eroded, cv.MORPH_CLOSE, horizontal_kernel)
    vertical_kernel = np.ones((int(height * kernel_length), 1))
    vertical = cv.morphologyEx(eroded, cv.MORPH_CLOSE, vertical_kernel)
    lines = cv.bitwise_and(vertical, horizontal)
    lines = cv.erode(lines, np.ones((3, 3)), iterations=3)
    return lines


def find_borders(coord_list) -> List[int]:
    """
    :param coord_list: list with border coords
    :return: list
    """
    return [0] + [index + 1 for index in range(len(coord_list) - 1)
                  if coord_list[index + 1] - coord_list[index] > 10]


def find_horizontal(inv_bin_image: np.array) -> List[int]:
    """
    :param inv_bin_image: inverted image with binarization
    :return:
    """
    return [index for index, line in enumerate(inv_bin_image)
            if np.count_nonzero(line) >= inv_bin_image.shape[1] - inv_bin_image.shape[1] * 0.0375]


def find_vertical(inv_bin_image: np.array) -> List[int]:
    """
    :param inv_bin_image: inverted image with binarization
    :return:
    """
    return [index for index in range(inv_bin_image.shape[1])
            if np.count_nonzero(inv_bin_image[:, index]) >= inv_bin_image.shape[0] - inv_bin_image.shape[0] * 0.0375]


def check_on_table(inv_bin_image: np.array) -> bool:
    """
    :param inv_bin_image: inverted image with binarization
    :return: is there a table in the picture?
    """
    return not (len(find_horizontal(inv_bin_image)) == 0 or len(find_vertical(inv_bin_image)) == 0)


def convert_pdf_to_docx(pdf_file: bytes) -> None:
    """
    :param pdf_file: pdf file in bytesIo representation
    :return: void (create docx file in current dir)
    """
    document = Document()
    document.add_heading('PDF_TO_DOCX', 0)
    images = read_from_pdf(pdf_file)
    for img in images:
        cv.imwrite("static/temp.png", img)
        osd = pytesseract.image_to_osd("static/temp.png",
                                       config='--psm 0 -c min_characters_to_try=5')
        angle = int(re.search(r'(?<=Rotate: )\d+', osd).group(0))
        match angle:
            case 90:
                image_norm = cv.rotate(img, cv.ROTATE_90_CLOCKWISE)
            case 180:
                image_norm = cv.rotate(img, cv.ROTATE_180)
            case 270:
                img = cv.rotate(img, cv.ROTATE_180)
                image_norm = cv.rotate(img, cv.ROTATE_90_CLOCKWISE)
            case 360:
                img = cv.rotate(img, cv.ROTATE_180)
                image_norm = cv.rotate(img, cv.ROTATE_180)
            case _:
                image_norm = img

        cropped_parts = get_cropped_parts(image_norm)
        for crop in cropped_parts:
            gray = cv.cvtColor(crop, cv.COLOR_BGR2GRAY)
            lines = find_lines(gray)
            _, inv_bin_image = cv.threshold(lines, 200, 255, cv.THRESH_BINARY_INV)

            if check_on_table(inv_bin_image):
                horizontal = find_horizontal(inv_bin_image)
                vertical = find_vertical(inv_bin_image)
                h_borders = find_borders(horizontal)
                v_borders = find_borders(vertical)

                table = document.add_table(rows=len(h_borders) - 1,
                                           cols=len(v_borders) - 1,
                                           style='Table Grid')

                for row in range(len(v_borders) - 1):
                    for col in range(len(h_borders) - 1):
                        cr = crop[horizontal[h_borders[col]]:horizontal[h_borders[col + 1]],
                                  vertical[v_borders[row]]:vertical[v_borders[row + 1]]]
                        text = pytesseract.image_to_string(cr, lang="rus").strip()
                        try:
                            column = table.columns[row].cells
                            column[col].text = text
                        except Exception as ex:
                            print("Something wrong with symbols recognition")
                            print(ex.with_traceback(None))

            else:
                text = pytesseract.image_to_string(crop, lang="rus").strip()
                if len(re.sub(r"\s+", "", text)) == 0:
                    cv.imwrite("static/temp.png", crop)
                    document.add_picture('static/temp.png')
                else:
                    try:
                        document.add_paragraph(text)
                    except Exception as ex:
                        print("Something wrong with symbols recognition")
                        print(ex.with_traceback(None))

    os.remove("/static/temp.png")
    document.save('PDF_TO_DOCX.docx')
